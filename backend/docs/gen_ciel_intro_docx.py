"""Generate the Ciel / chatRAG introduction document (bilingual)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "Ciel_Introduction.docx"

ACCENT = RGBColor(0x3B, 0x82, 0xF6)


def h(doc, text, size=16):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = ACCENT
    return p


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


doc = Document()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Ciel — Trợ lý AI của chatRAG")
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Tài liệu giới thiệu chính thức · Official Introduction")
sr.italic = True
sr.font.size = Pt(11)

doc.add_paragraph()

# Vietnamese section
h(doc, "🇻🇳 Tiếng Việt", 18)
doc.add_paragraph(
    "Xin chào! Tôi là Ciel — trợ lý AI của chatRAG. Tôi giúp bạn tra cứu và "
    "hỏi đáp trực tiếp trên kho tài liệu nội bộ của bạn, luôn trả lời kèm "
    "trích dẫn nguồn để bạn dễ kiểm chứng."
)

h(doc, "Tôi làm được gì", 13)
for t in [
    "Trả lời câu hỏi dựa trên tài liệu bạn đã tải lên (RAG) — kèm nguồn trích dẫn.",
    "Đọc nhiều định dạng: PDF, ảnh (OCR), Word (.docx), Excel (.xlsx).",
    "Tổ chức tài liệu theo thư mục; giới hạn phạm vi tìm kiếm theo từng thư mục.",
    "Chế độ Hybrid: kết hợp tài liệu với kiến thức chung khi tài liệu chưa đủ.",
    "Đa ngôn ngữ: hiểu và trả lời tiếng Việt, tiếng Anh, tiếng Nhật...",
    "Nhiều mô hình AI: chạy local qua Ollama, hoặc cloud tốc độ cao qua Groq.",
]:
    bullet(doc, t)

h(doc, "Cách dùng nhanh", 13)
for t in [
    'Tải tài liệu lên ở "Upload files" hoặc "Sync folder".',
    "Gõ câu hỏi vào ô chat.",
    'Chọn thư mục ở thanh dưới để giới hạn phạm vi (hoặc "All").',
    "Gõ /help bất cứ lúc nào để xem lại phần giới thiệu này.",
]:
    numbered(doc, t)

doc.add_paragraph()

# English section
h(doc, "🇬🇧 English", 18)
doc.add_paragraph(
    "Hello! I'm Ciel — the AI assistant of chatRAG. I help you search and ask "
    "questions directly over your internal knowledge base, always answering "
    "with source citations so you can verify everything."
)

h(doc, "What I can do", 13)
for t in [
    "Answer questions grounded in your uploaded documents (RAG) — with cited sources.",
    "Read many formats: PDF, images (OCR), Word (.docx), Excel (.xlsx).",
    "Organize documents into folders; scope the search to specific folders.",
    "Hybrid mode: blend your documents with general knowledge when docs fall short.",
    "Multilingual: I understand and reply in Vietnamese, English, Japanese, and more.",
    "Multiple AI models: run locally via Ollama, or fast cloud models via Groq.",
]:
    bullet(doc, t)

h(doc, "Quick start", 13)
for t in [
    'Upload documents via "Upload files" or "Sync folder".',
    "Type your question in the chat box.",
    'Pick a folder in the bottom bar to scope the search (or "All").',
    "Type /help anytime to see this introduction again.",
]:
    numbered(doc, t)

doc.save(OUT)
print(f"Saved: {OUT}")
