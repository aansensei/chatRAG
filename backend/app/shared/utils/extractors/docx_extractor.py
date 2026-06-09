import zipfile
from pathlib import Path
from docx import Document
from .base import ExtractResult


def extract_docx(file_path: str, image_output_dir: str = "extracted_images") -> ExtractResult:
    path = Path(file_path)
    image_dir = Path(image_output_dir)
    image_dir.mkdir(parents=True, exist_ok=True)

    result = ExtractResult()
    doc = Document(file_path)
    text_parts = []
    tables = []
    image_paths = []

    # text: iterate paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # tables: first row as header, remaining rows as list of dicts
    # row.cells repeats merged cells — deduplicate by XML element identity
    def _unique_cells(row):
        seen = set()
        cells = []
        for cell in row.cells:
            if id(cell._tc) not in seen:
                seen.add(id(cell._tc))
                cells.append(cell.text.strip())
        return cells

    for table_index, table in enumerate(doc.tables):
        rows = [_unique_cells(row) for row in table.rows]
        if not rows:
            continue
        header = rows[0]
        data = [
            {header[i]: cell for i, cell in enumerate(row)}
            for row in rows[1:]
        ]
        tables.append({
            "table_index": table_index + 1,
            "data": data
        })

    # images: a docx file is a zip archive, images live under word/media/
    with zipfile.ZipFile(file_path, "r") as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                filename = Path(name).name
                save_path = image_dir / f"{path.stem}_{filename}"
                save_path.write_bytes(z.read(name))
                image_paths.append(str(save_path))

    result.text = "\n".join(text_parts)
    result.tables = tables
    result.images = image_paths
    result.metadata = {"source": str(path)}

    return result
