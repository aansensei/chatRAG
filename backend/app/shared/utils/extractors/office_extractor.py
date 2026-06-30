import csv
from pathlib import Path
from .base import ExtractResult


def extract_csv(file_path: str) -> ExtractResult:
    rows = []
    for enc in ("utf-8-sig", "utf-8", "cp1258", "cp1252", "latin-1"):
        try:
            with open(file_path, newline="", encoding=enc) as f:
                reader = csv.reader(f)
                for row in reader:
                    rows.append("  |  ".join(cell.strip() for cell in row))
            break
        except (UnicodeDecodeError, Exception):
            rows = []
    result = ExtractResult()
    result.text = "\n".join(rows)
    result.metadata = {"source": str(Path(file_path).name), "pages": 1}
    return result


def extract_docx(file_path: str) -> ExtractResult:
    from docx import Document
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append("  |  ".join(cells))
        if rows:
            parts.append("\n".join(rows))
    result = ExtractResult()
    result.text = "\n\n".join(parts)
    result.metadata = {"source": str(Path(file_path).name), "pages": 1}
    return result


def extract_xlsx(file_path: str) -> ExtractResult:
    import openpyxl
    wb = openpyxl.load_workbook(file_path, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = "  |  ".join(cells).strip(" |")
            if line:
                rows.append(line)
        if rows:
            parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
    result = ExtractResult()
    result.text = "\n\n".join(parts)
    result.metadata = {"source": str(Path(file_path).name), "pages": len(wb.sheetnames)}
    return result
